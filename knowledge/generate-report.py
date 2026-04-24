#!/usr/bin/env python3
"""
Generate a formatted DOCX synthetic user test report.

Usage:
    python3 generate-report.py output.docx

This script reads report data from stdin as JSON and produces a
professionally formatted DOCX file with:
- Landscape orientation
- Smaller font sizes
- Key:value metadata on separate lines
- Bold highlights in summary
- Priority tables grouped by P-level headers (not as column values)
- Clean table formatting

Input JSON format:
{
    "date": "2026-04-22",
    "test_mode": "Comparative (As-Is → To-Be)",
    "personas": "Anna Lee — Senior Game Developer (Tech 9/10), Alex Johnson — Junior Game Developer (Tech 5/10)",
    "summary": "The **redesign resolves 3 of 5** critical As-Is complaints...",
    "fixes": [
        {"priority": "P0", "issue": "...", "heuristic": "H3", "personas": "Both", "recommendation": "..."},
        ...
    ],
    "real_user_testing": [
        {"area": "...", "why": "...", "method": "..."},
        ...
    ]
}
"""

import json
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex
    })
    shading.append(shading_elem)


def add_bold_text(paragraph, text, font_size=9):
    """Add text with **bold** markdown markers converted to actual bold."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'


def set_table_font(table, size=8):
    """Set font size for all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size)
                    run.font.name = 'Arial'


def generate_report(data, output_path):
    doc = Document()

    # --- Landscape orientation ---
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # --- Title ---
    title = doc.add_heading('Synthetic User Test Report', level=1)
    title.runs[0].font.size = Pt(16)

    # --- Metadata block (each on new line) ---
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(4)

    run = meta.add_run('Date: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run = meta.add_run(data['date'] + '\n')
    run.font.size = Pt(9)
    run.font.name = 'Arial'

    run = meta.add_run('Test mode: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run = meta.add_run(data['test_mode'] + '\n')
    run.font.size = Pt(9)
    run.font.name = 'Arial'

    run = meta.add_run('Personas: ')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run = meta.add_run(data['personas'])
    run.font.size = Pt(9)
    run.font.name = 'Arial'

    # --- Summary ---
    doc.add_heading('Summary', level=2).runs[0].font.size = Pt(12)
    summary_para = doc.add_paragraph()
    summary_para.paragraph_format.space_after = Pt(8)
    add_bold_text(summary_para, data['summary'], font_size=9)

    # --- Priority Fixes ---
    doc.add_heading('Priority Fixes', level=2).runs[0].font.size = Pt(12)

    # Group fixes by priority
    priorities = {}
    for fix in data['fixes']:
        p = fix['priority']
        if p not in priorities:
            priorities[p] = []
        priorities[p].append(fix)

    priority_colors = {
        'P0': 'FF4444',  # Red
        'P1': 'FF8800',  # Orange
        'P2': 'FFBB00',  # Yellow
        'P3': 'AAAAAA',  # Gray
    }

    priority_labels = {
        'P0': 'P0 — Critical (fix immediately)',
        'P1': 'P1 — Major (fix before release)',
        'P2': 'P2 — Minor (plan for next sprint)',
        'P3': 'P3 — Cosmetic (backlog)',
    }

    for priority in ['P0', 'P1', 'P2', 'P3']:
        if priority not in priorities:
            continue

        fixes = priorities[priority]

        # Priority header
        p_header = doc.add_paragraph()
        run = p_header.add_run(priority_labels.get(priority, priority))
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(
            int(priority_colors[priority][:2], 16),
            int(priority_colors[priority][2:4], 16),
            int(priority_colors[priority][4:6], 16)
        )

        # Table for this priority group
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Header row — Issue, Recommendation, Heuristic, Personas
        headers = ['Issue', 'Recommendation', 'Heuristic', 'Personas']
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = ''
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
            run.font.size = Pt(8)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, '333333')

        # Data rows
        for fix in fixes:
            row = table.add_row()
            cells = row.cells
            cells[0].text = fix['issue']
            cells[1].text = fix['recommendation']
            cells[2].text = fix['heuristic']
            cells[3].text = fix['personas']

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Cm(7)
            row.cells[1].width = Cm(11)
            row.cells[2].width = Cm(3)
            row.cells[3].width = Cm(3)

        set_table_font(table, size=8)
        doc.add_paragraph('')  # Spacer

    # --- Real-User Testing ---
    doc.add_heading('Real-User Testing Recommendations', level=2).runs[0].font.size = Pt(12)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers = ['#', 'What to test with real users', 'Why synthetic testing can\'t validate', 'Suggested method']
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '333333')

    for i, item in enumerate(data['real_user_testing']):
        row = table.add_row()
        cells = row.cells
        cells[0].text = str(i + 1)
        cells[1].text = item['area']
        cells[2].text = item['why']
        cells[3].text = item['method']

    for row in table.rows:
        row.cells[0].width = Cm(0.3)
        row.cells[1].width = Cm(8.5)
        row.cells[2].width = Cm(8.5)
        row.cells[3].width = Cm(6.7)
        # Center-align the # column
        for paragraph in row.cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_table_font(table, size=8)

    # --- Disclaimer ---
    doc.add_paragraph('')
    disclaimer = doc.add_paragraph()
    run = disclaimer.add_run('These findings are hypotheses generated by synthetic personas — not validated results from real users. Use as input for your research plan.')
    run.italic = True
    run.font.size = Pt(7)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python3 generate-report.py output.docx < report.json")
        sys.exit(1)

    data = json.load(sys.stdin)
    output = generate_report(data, sys.argv[1])
    print(f"Report generated: {output}")
